
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _find_heading(doc, text):
    for p in doc.paragraphs:
        if p.text.strip().lower() == text.strip().lower():
            return p
    return None

def _append_heading(doc, title, level=2):
    # Create a heading at end, return the paragraph object
    return doc.add_heading(title, level=level)

def _move_paragraph_after(anchor_paragraph, paragraph_to_move):
    # Move an existing paragraph so it appears immediately after the anchor paragraph
    anchor_paragraph._element.addnext(paragraph_to_move._p)

def _move_table_after(anchor_paragraph, table_to_move):
    # Move an existing table so it appears immediately after the anchor paragraph
    anchor_paragraph._element.addnext(table_to_move._tbl)

def add_table_after(doc, anchor_text, df, title=None):
    anchor = _find_heading(doc, anchor_text)
    if anchor is None:
        # Fallback: append at end
        if title:
            doc.add_heading(title, level=2)
        table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
    else:
        # Optionally insert a subheading first
        if title:
            h = _append_heading(doc, title, level=2)
            _move_paragraph_after(anchor, h)
            anchor = h  # insert the table after the new subheading
        table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
        _move_table_after(anchor, table)

    # header
    for j, col in enumerate(df.columns):
        table.cell(0, j).text = str(col)
    # body
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            val = "" if df.iat[i, j] is None else str(df.iat[i, j])
            table.cell(i+1, j).text = val
    return table

def add_paragraph_after(doc, anchor_text, content, title=None):
    anchor = _find_heading(doc, anchor_text)
    if anchor is None:
        if title:
            doc.add_heading(title, level=2)
        p = doc.add_paragraph(content)
        return

    # Insert optional subheading
    if title:
        h = _append_heading(doc, title, level=2)
        _move_paragraph_after(anchor, h)
        anchor = h

    # Now add the main paragraph and move it after the (sub)heading
    p = doc.add_paragraph(content)
    _move_paragraph_after(anchor, p)

def save_docx_with_sections(template_path, output_path, replacements, sections):
    """
    sections: list of dicts with keys:
      - kind: 'table' or 'text'
      - anchor: heading text from template where content should be placed
      - title: optional subheading (level-2 style)
      - content: DataFrame (for table) or str (for text)
    """
    doc = Document(template_path)

    # simple replacements in whole document text (title page)
    for p in doc.paragraphs:
        for key, val in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, str(val))

    for sec in sections:
        kind = sec.get("kind")
        anchor = sec.get("anchor")
        title = sec.get("title")
        content = sec.get("content")
        if kind == "table":
            add_table_after(doc, anchor, content, title=title)
        elif kind == "text":
            add_paragraph_after(doc, anchor, content, title=title)

    doc.save(output_path)
